#!/usr/bin/env python3
"""
Brand Deep Research Word Document Generator

Converts a JSON brand deep research profile (following the schema in
schemas/brand_deep_research_schema.md) into a branded Upscale Word
document (.docx) with seven sections:

1. Cover / Title Page
2. Company Overview
3. Products & Services
4. Performance
5. Outlook & Strategy
6. Brand Profile
7. Research Notes & Citations

Uses Upscale brand colours:
- Eerie Black: #191919 (headings, body text)
- Lime Green: #34C52A (section labels, accents, rules)
- Dark Cyan: #429792 (metadata, captions)
- Ivory: #FEFFEA (callout backgrounds)

Font: Lexend throughout
"""

import json
import sys
import os
from datetime import datetime

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# ---------------------------------------------------------------------------
# Upscale Brand Colours
# ---------------------------------------------------------------------------
EERIE_BLACK = RGBColor(0x19, 0x19, 0x19)
LIME_GREEN = RGBColor(0x34, 0xC5, 0x2A)
DARK_CYAN = RGBColor(0x42, 0x97, 0x92)
IVORY = RGBColor(0xFE, 0xFF, 0xEA)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY_MID = RGBColor(0x66, 0x66, 0x66)

FONT_NAME = "Lexend"


# ---------------------------------------------------------------------------
# Helper: set cell shading
# ---------------------------------------------------------------------------
def set_cell_shading(cell, hex_color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


# ---------------------------------------------------------------------------
# Helper: styled paragraph
# ---------------------------------------------------------------------------
def add_styled_paragraph(doc, text, font_size=10, bold=False, color=None,
                         alignment=None, space_after=6, space_before=0,
                         italic=False, all_caps=False):
    """Add a paragraph with Upscale styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color or EERIE_BLACK
    run.font.all_caps = all_caps
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if alignment:
        p.alignment = alignment
    return p


# ---------------------------------------------------------------------------
# Helper: section label (uppercase, Lime Green, spaced)
# ---------------------------------------------------------------------------
def add_section_label(doc, text):
    """Add an uppercase Lime Green section label."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.font.name = FONT_NAME
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = LIME_GREEN
    run.font.all_caps = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    return p


# ---------------------------------------------------------------------------
# Helper: heading
# ---------------------------------------------------------------------------
def add_heading_text(doc, text, level=1):
    """Add a heading with Upscale styling."""
    sizes = {1: 20, 2: 16, 3: 13}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(sizes.get(level, 13))
    run.font.bold = True
    run.font.color.rgb = EERIE_BLACK
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    return p


# ---------------------------------------------------------------------------
# Helper: arrow bullet
# ---------------------------------------------------------------------------
def add_arrow_bullet(doc, text, font_size=10, color=None):
    """Add a line with arrow bullet prefix."""
    p = doc.add_paragraph()
    arrow_run = p.add_run("→  ")
    arrow_run.font.name = FONT_NAME
    arrow_run.font.size = Pt(font_size)
    arrow_run.font.color.rgb = LIME_GREEN
    arrow_run.font.bold = True
    text_run = p.add_run(text)
    text_run.font.name = FONT_NAME
    text_run.font.size = Pt(font_size)
    text_run.font.color.rgb = color or EERIE_BLACK
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.3)
    return p


# ---------------------------------------------------------------------------
# Helper: horizontal rule (Lime Green)
# ---------------------------------------------------------------------------
def add_horizontal_rule(doc):
    """Add a Lime Green horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Use a bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="34C52A"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return p


# ---------------------------------------------------------------------------
# Helper: key-value line
# ---------------------------------------------------------------------------
def add_key_value(doc, key, value, key_color=None):
    """Add a key: value line."""
    p = doc.add_paragraph()
    key_run = p.add_run(f"{key}:  ")
    key_run.font.name = FONT_NAME
    key_run.font.size = Pt(10)
    key_run.font.bold = True
    key_run.font.color.rgb = key_color or DARK_CYAN
    val_run = p.add_run(str(value))
    val_run.font.name = FONT_NAME
    val_run.font.size = Pt(10)
    val_run.font.color.rgb = EERIE_BLACK
    p.paragraph_format.space_after = Pt(3)
    return p


# ---------------------------------------------------------------------------
# Helper: branded table
# ---------------------------------------------------------------------------
def add_branded_table(doc, headers, rows, col_widths=None):
    """Add a table with Eerie Black header row and Ivory data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_NAME
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = WHITE
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, "191919")

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text) if cell_text else "—")
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.color.rgb = EERIE_BLACK
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if row_idx % 2 == 0:
                set_cell_shading(cell, "FEFFEA")

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    # Add spacing after table
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    return table


# ---------------------------------------------------------------------------
# Helper: citation reference (superscript)
# ---------------------------------------------------------------------------
def add_text_with_citation(doc, text, citation_id=None, font_size=10):
    """Add text with optional superscript citation number."""
    p = doc.add_paragraph()
    text_run = p.add_run(text)
    text_run.font.name = FONT_NAME
    text_run.font.size = Pt(font_size)
    text_run.font.color.rgb = EERIE_BLACK
    if citation_id:
        cite_run = p.add_run(f" [{citation_id}]")
        cite_run.font.name = FONT_NAME
        cite_run.font.size = Pt(7)
        cite_run.font.color.rgb = DARK_CYAN
        cite_run.font.superscript = True
    p.paragraph_format.space_after = Pt(4)
    return p


# ---------------------------------------------------------------------------
# Helper: safe get with fallback
# ---------------------------------------------------------------------------
def safe_get(data, *keys, default="—"):
    """Safely traverse nested dict keys."""
    current = data
    for key in keys:
        if isinstance(current, dict):
            current = current.get(key, default)
        else:
            return default
    return current if current is not None else default


# ===================================================================
# SECTION BUILDERS
# ===================================================================

def build_cover_page(doc, data):
    """Build the cover / title page."""
    # Company name as hero heading
    p = doc.add_paragraph()
    run = p.add_run(data.get("company_name", "Brand Research"))
    run.font.name = FONT_NAME
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = EERIE_BLACK
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(4)

    # Subtitle
    add_styled_paragraph(
        doc,
        "Brand Deep Research Profile",
        font_size=14, color=LIME_GREEN, bold=True,
        space_after=12
    )

    add_horizontal_rule(doc)

    # Metadata
    add_key_value(doc, "Industry", data.get("industry", "—"))
    add_key_value(doc, "Website", data.get("website", "—"))
    add_key_value(doc, "Analysis Date", data.get("analysis_date", "—")[:10])
    add_key_value(doc, "Research Depth",
                  data.get("research_depth", "—").replace("_", " ").title())
    add_key_value(doc, "Profile Completeness",
                  f"{data.get('overall_profile_completeness', '—')} / 10")

    add_horizontal_rule(doc)

    # Executive summary
    add_section_label(doc, "Executive Summary")
    add_styled_paragraph(
        doc,
        data.get("executive_summary", "No executive summary provided."),
        font_size=11, space_after=12
    )

    # Key gaps
    gaps = data.get("key_gaps", [])
    if gaps:
        add_section_label(doc, "Key Research Gaps")
        for gap in gaps:
            add_arrow_bullet(doc, gap, color=GREY_MID)

    doc.add_page_break()


def build_company_overview(doc, data):
    """Build the Company Overview section."""
    overview = data.get("company_overview", {})

    add_section_label(doc, "Company Overview")
    add_styled_paragraph(doc, overview.get("summary", ""), font_size=10,
                         space_after=8)
    add_horizontal_rule(doc)

    # Key facts grid
    add_heading_text(doc, "Key Facts", level=2)
    facts = [
        ("Legal Name", overview.get("legal_name", "—")),
        ("Trading As", overview.get("trading_as", "—")),
        ("Founded", overview.get("founded", "—")),
        ("Headquarters", overview.get("headquarters", "—")),
        ("Company Type", safe_get(overview, "company_type", default="—").replace("_", " ").title()),
        ("Stock Ticker", overview.get("stock_ticker") or "N/A"),
        ("Employees", overview.get("employee_count", "—")),
        ("Employee Trend", safe_get(overview, "employee_count_trend", default="—").replace("_", " ").title()),
        ("Annual Revenue", overview.get("annual_revenue", "—")),
        ("Ownership", overview.get("ownership_structure", "—")),
    ]
    for key, value in facts:
        add_key_value(doc, key, value)

    # Mission & Vision
    add_heading_text(doc, "Mission & Vision", level=2)
    add_key_value(doc, "Mission", overview.get("mission_statement", "—"))
    vision = overview.get("vision_statement")
    if vision:
        add_key_value(doc, "Vision", vision)

    # Core Values
    values = overview.get("core_values", [])
    if values:
        add_heading_text(doc, "Core Values", level=3)
        for v in values:
            add_arrow_bullet(doc, v)

    # Leadership
    leaders = overview.get("key_leadership", [])
    if leaders:
        add_heading_text(doc, "Key Leadership", level=2)
        headers = ["Name", "Title", "Tenure", "Background"]
        rows = []
        for person in leaders:
            rows.append([
                person.get("name", "—"),
                person.get("title", "—"),
                person.get("tenure", "—"),
                person.get("background", "—")[:120],
            ])
        add_branded_table(doc, headers, rows, col_widths=[1.5, 1.8, 1.2, 2.5])

    # Major Shareholders
    shareholders = overview.get("major_shareholders", [])
    if shareholders:
        add_heading_text(doc, "Major Shareholders", level=3)
        for sh in shareholders:
            add_arrow_bullet(doc, sh)

    # Geographic Presence
    geo = overview.get("geographic_presence", [])
    if geo:
        add_heading_text(doc, "Geographic Presence", level=3)
        for g in geo:
            add_arrow_bullet(doc, g)

    # Key Milestones
    milestones = overview.get("key_milestones", [])
    if milestones:
        add_heading_text(doc, "Key Milestones", level=2)
        headers = ["Date", "Milestone", "Source", "Confidence"]
        rows = []
        for m in milestones:
            rows.append([
                m.get("date", "—"),
                m.get("fact", "—")[:120],
                m.get("source_name", "—"),
                m.get("confidence", "—").upper(),
            ])
        add_branded_table(doc, headers, rows, col_widths=[1.0, 3.5, 1.5, 1.0])

    # Research confidence
    conf = overview.get("research_confidence", "—")
    add_styled_paragraph(
        doc,
        f"Section Research Confidence: {conf.upper()}",
        font_size=9, italic=True, color=DARK_CYAN, space_before=8
    )

    doc.add_page_break()


def build_products_services(doc, data):
    """Build the Products & Services section."""
    products = data.get("products_services", {})

    add_section_label(doc, "Products & Services")
    add_styled_paragraph(doc, products.get("summary", ""), font_size=10,
                         space_after=8)
    add_horizontal_rule(doc)

    # Product Portfolio
    portfolio = products.get("product_portfolio", [])
    if portfolio:
        add_heading_text(doc, "Product Portfolio", level=2)
        headers = ["Product", "Category", "Position", "Audience", "Pricing",
                   "Revenue %"]
        rows = []
        for p in portfolio[:20]:  # Cap at 20
            rows.append([
                p.get("name", "—"),
                p.get("category", "—"),
                safe_get(p, "market_position", default="—").replace("_", " ").title(),
                p.get("target_audience", "—")[:60],
                safe_get(p, "pricing_tier", default="—").replace("_", " ").title(),
                p.get("estimated_revenue_contribution", "—"),
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[1.3, 1.3, 1.0, 1.5, 0.9, 1.0])

        # Key differentiators for each product
        for p in portfolio[:20]:
            diffs = p.get("key_differentiators", [])
            if diffs:
                add_heading_text(doc, f"{p.get('name', '—')} — Key Differentiators",
                                 level=3)
                for d in diffs:
                    add_arrow_bullet(doc, d)

    # Service Offerings
    services = products.get("service_offerings", [])
    if services:
        add_heading_text(doc, "Service Offerings", level=2)
        headers = ["Service", "Category", "Audience", "Pricing", "Revenue %"]
        rows = []
        for s in services:
            rows.append([
                s.get("name", "—"),
                s.get("category", "—"),
                s.get("target_audience", "—")[:60],
                safe_get(s, "pricing_tier", default="—").replace("_", " ").title(),
                s.get("estimated_revenue_contribution", "—"),
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[1.5, 1.3, 2.0, 1.0, 1.2])

    # Flagship Products
    flagship = products.get("flagship_products", [])
    if flagship:
        add_heading_text(doc, "Flagship Products", level=3)
        for f in flagship:
            add_arrow_bullet(doc, f)

    # Recent Launches
    launches = products.get("recent_launches", [])
    if launches:
        add_heading_text(doc, "Recent Launches", level=2)
        for launch in launches:
            add_text_with_citation(
                doc,
                f"{launch.get('date', '—')}: {launch.get('fact', '—')}",
            )
            add_styled_paragraph(
                doc,
                f"Source: {launch.get('source_name', '—')}",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Discontinued Products
    discontinued = products.get("discontinued_products", [])
    if discontinued:
        add_heading_text(doc, "Discontinued Products", level=3)
        for d in discontinued:
            add_arrow_bullet(doc, f"{d.get('date', '—')}: {d.get('fact', '—')}")

    # R&D Focus
    rnd = products.get("r_and_d_focus", "")
    if rnd:
        add_heading_text(doc, "R&D Focus", level=3)
        add_styled_paragraph(doc, rnd)

    # IP
    ip = products.get("intellectual_property", "")
    if ip:
        add_heading_text(doc, "Intellectual Property", level=3)
        add_styled_paragraph(doc, ip)

    # Lifecycle
    lifecycle = products.get("product_lifecycle_stage", "")
    if lifecycle:
        add_key_value(doc, "Product Lifecycle Stage",
                      lifecycle.replace("_", " ").title())

    # Confidence
    conf = products.get("research_confidence", "—")
    add_styled_paragraph(
        doc,
        f"Section Research Confidence: {conf.upper()}",
        font_size=9, italic=True, color=DARK_CYAN, space_before=8
    )

    doc.add_page_break()


def build_performance(doc, data):
    """Build the Performance section."""
    perf = data.get("performance_profile", {})

    add_section_label(doc, "Performance")
    add_styled_paragraph(doc, perf.get("summary", ""), font_size=10,
                         space_after=8)
    add_horizontal_rule(doc)

    # Financial Metrics
    metrics = perf.get("financial_metrics", [])
    if metrics:
        add_heading_text(doc, "Financial Metrics", level=2)
        headers = ["Metric", "Value", "Period", "Trend", "Source"]
        rows = []
        for m in metrics:
            rows.append([
                m.get("metric", "—"),
                m.get("value", "—"),
                m.get("period", "—"),
                m.get("trend", "—").replace("_", " ").title(),
                m.get("source_name", "—"),
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[1.5, 1.0, 1.5, 1.0, 2.0])

    # Summary metrics
    add_heading_text(doc, "Performance Assessment", level=2)
    add_key_value(doc, "Revenue Trend",
                  safe_get(perf, "revenue_trend", default="—").replace("_", " ").title())
    add_key_value(doc, "Profitability",
                  safe_get(perf, "profitability_assessment", default="—").replace("_", " ").title())
    add_key_value(doc, "Market Share", perf.get("market_share", "—"))
    add_key_value(doc, "Market Share Trend",
                  safe_get(perf, "market_share_trend", default="—").replace("_", " ").title())
    add_key_value(doc, "Growth Rate", perf.get("growth_rate", "—"))
    add_key_value(doc, "Competitive Ranking",
                  perf.get("competitive_ranking", "—"))

    # KPIs
    kpis = perf.get("key_kpis", [])
    if kpis:
        add_heading_text(doc, "Key Performance Indicators", level=2)
        headers = ["KPI", "Value", "Benchmark", "Assessment"]
        rows = []
        for k in kpis:
            rows.append([
                k.get("kpi_name", "—"),
                k.get("value", "—"),
                k.get("benchmark") or "—",
                k.get("assessment", "—")[:80],
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[1.5, 1.0, 1.5, 3.0])

    # Highlights
    highlights = perf.get("recent_performance_highlights", [])
    if highlights:
        add_heading_text(doc, "Recent Performance Highlights", level=2)
        for h in highlights:
            add_text_with_citation(doc, h.get("fact", "—"))
            add_styled_paragraph(
                doc,
                f"Source: {h.get('source_name', '—')} ({h.get('date', '—')})",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Concerns
    concerns = perf.get("recent_performance_concerns", [])
    if concerns:
        add_heading_text(doc, "Performance Concerns", level=2)
        for c in concerns:
            add_text_with_citation(doc, c.get("fact", "—"))
            add_styled_paragraph(
                doc,
                f"Source: {c.get('source_name', '—')} ({c.get('date', '—')})",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Confidence
    conf = perf.get("research_confidence", "—")
    add_styled_paragraph(
        doc,
        f"Section Research Confidence: {conf.upper()}",
        font_size=9, italic=True, color=DARK_CYAN, space_before=8
    )

    doc.add_page_break()


def build_outlook_strategy(doc, data):
    """Build the Outlook & Strategy section."""
    outlook = data.get("outlook_strategy", {})

    add_section_label(doc, "Outlook & Strategy")
    add_styled_paragraph(doc, outlook.get("summary", ""), font_size=10,
                         space_after=8)
    add_horizontal_rule(doc)

    # Stated Strategy
    strategy = outlook.get("stated_strategy", "")
    if strategy:
        add_heading_text(doc, "Stated Strategy", level=2)
        add_styled_paragraph(doc, strategy)

    # Strategic Initiatives
    initiatives = outlook.get("strategic_initiatives", [])
    if initiatives:
        add_heading_text(doc, "Strategic Initiatives", level=2)
        headers = ["Initiative", "Status", "Timeline", "Rationale", "Source"]
        rows = []
        for i in initiatives:
            rows.append([
                i.get("initiative", "—")[:80],
                i.get("status", "—").replace("_", " ").title(),
                i.get("timeline", "—"),
                i.get("strategic_rationale", "—")[:80],
                i.get("source_name", "—"),
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[2.0, 0.8, 1.2, 2.0, 1.0])

    # Growth Plans
    growth = outlook.get("growth_plans", [])
    if growth:
        add_heading_text(doc, "Growth Plans", level=2)
        for g in growth:
            add_text_with_citation(doc, g.get("fact", "—"))
            add_styled_paragraph(
                doc,
                f"Source: {g.get('source_name', '—')} | Confidence: {g.get('confidence', '—').upper()}",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Expansion Targets
    targets = outlook.get("expansion_targets", [])
    if targets:
        add_heading_text(doc, "Expansion Targets", level=3)
        for t in targets:
            add_arrow_bullet(doc, t)

    # Investment Priorities
    priorities = outlook.get("investment_priorities", [])
    if priorities:
        add_heading_text(doc, "Investment Priorities", level=3)
        for p in priorities:
            add_arrow_bullet(doc, p)

    # Industry Trends
    trends = outlook.get("industry_trends_affecting", [])
    if trends:
        add_heading_text(doc, "Industry Trends Affecting", level=2)
        headers = ["Trend", "Impact", "Timeframe"]
        rows = []
        for t in trends:
            rows.append([
                t.get("trend", "—")[:100],
                t.get("impact", "—").replace("_", " ").title(),
                t.get("timeframe", "—").replace("_", " ").title(),
            ])
        add_branded_table(doc, headers, rows, col_widths=[4.0, 1.0, 1.0])

    # Key Risks
    risks = outlook.get("key_risks", [])
    if risks:
        add_heading_text(doc, "Key Risks", level=2)
        headers = ["Risk", "Severity", "Likelihood", "Mitigation"]
        rows = []
        for r in risks:
            rows.append([
                r.get("risk", "—")[:100],
                r.get("severity", "—").upper(),
                r.get("likelihood", "—").upper(),
                (r.get("mitigation") or "—")[:80],
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[2.5, 1.0, 1.0, 2.5])

    # Analyst Consensus
    consensus = outlook.get("analyst_consensus", "")
    if consensus:
        add_heading_text(doc, "Analyst Consensus", level=3)
        add_styled_paragraph(doc, consensus)

    # Growth Outlook
    growth_outlook = outlook.get("growth_outlook", "")
    if growth_outlook:
        add_key_value(doc, "Growth Outlook",
                      growth_outlook.replace("_", " ").title())

    # Confidence
    conf = outlook.get("research_confidence", "—")
    add_styled_paragraph(
        doc,
        f"Section Research Confidence: {conf.upper()}",
        font_size=9, italic=True, color=DARK_CYAN, space_before=8
    )

    doc.add_page_break()


def build_brand_profile(doc, data):
    """Build the Brand Profile section."""
    brand = data.get("brand_profile", {})

    add_section_label(doc, "Brand Profile")
    add_styled_paragraph(doc, brand.get("summary", ""), font_size=10,
                         space_after=8)
    add_horizontal_rule(doc)

    # Brand Identity
    add_heading_text(doc, "Brand Identity", level=2)
    add_key_value(doc, "Positioning", brand.get("brand_positioning", "—"))
    add_key_value(doc, "Brand Promise", brand.get("brand_promise", "—"))
    archetype = brand.get("brand_archetype")
    if archetype:
        add_key_value(doc, "Brand Archetype", archetype)
    tagline = brand.get("tagline_slogan")
    if tagline:
        add_key_value(doc, "Tagline / Slogan", tagline)

    # Brand Voice & Visual Identity
    voice = brand.get("brand_voice_tone", "")
    if voice:
        add_heading_text(doc, "Brand Voice & Tone", level=3)
        add_styled_paragraph(doc, voice)

    visual = brand.get("visual_identity", "")
    if visual:
        add_heading_text(doc, "Visual Identity", level=3)
        add_styled_paragraph(doc, visual)

    # Target Audience
    audience = brand.get("target_audience", {})
    if audience:
        add_heading_text(doc, "Target Audience", level=2)
        add_key_value(doc, "Primary", audience.get("primary", "—"))
        secondary = audience.get("secondary")
        if secondary:
            add_key_value(doc, "Secondary", secondary)
        add_key_value(doc, "Demographics", audience.get("demographics", "—"))
        add_key_value(doc, "Psychographics",
                      audience.get("psychographics", "—"))

    # Messaging Pillars
    pillars = brand.get("key_messaging_pillars", [])
    if pillars:
        add_heading_text(doc, "Key Messaging Pillars", level=2)
        for p in pillars:
            add_arrow_bullet(doc, p)

    # Competitive Differentiation
    diffs = brand.get("competitive_differentiation", [])
    if diffs:
        add_heading_text(doc, "Competitive Differentiation", level=2)
        for d in diffs:
            add_arrow_bullet(doc, d)

    # Brand Perception
    perception = brand.get("brand_perception", {})
    if perception:
        add_heading_text(doc, "Brand Perception", level=2)
        add_key_value(doc, "Customer Sentiment",
                      safe_get(perception, "customer_sentiment", default="—").replace("_", " ").title())
        add_key_value(doc, "Industry Reputation",
                      safe_get(perception, "industry_reputation", default="—").replace("_", " ").title())
        add_key_value(doc, "Employer Brand",
                      safe_get(perception, "employer_brand", default="—").replace("_", " ").title())

        themes = perception.get("key_perception_themes", [])
        if themes:
            add_heading_text(doc, "Key Perception Themes", level=3)
            for t in themes:
                add_arrow_bullet(doc, t)

    # Competitors
    competitors = brand.get("competitors", [])
    if competitors:
        add_heading_text(doc, "Competitive Landscape", level=2)
        headers = ["Competitor", "Relationship", "Key Differentiator"]
        rows = []
        for c in competitors:
            rows.append([
                c.get("name", "—"),
                c.get("relationship", "—").replace("_", " ").title(),
                c.get("key_differentiator_vs_subject", "—")[:100],
            ])
        add_branded_table(doc, headers, rows,
                          col_widths=[1.5, 1.5, 4.0])

    # Brand Equity Indicators
    equity = brand.get("brand_equity_indicators", [])
    if equity:
        add_heading_text(doc, "Brand Equity Indicators", level=2)
        for e in equity:
            add_text_with_citation(doc, e.get("fact", "—"))
            add_styled_paragraph(
                doc,
                f"Source: {e.get('source_name', '—')}",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Digital Presence
    digital = brand.get("digital_presence", {})
    if digital:
        add_heading_text(doc, "Digital Presence", level=2)
        add_key_value(doc, "Website", digital.get("website", "—"))
        add_key_value(doc, "Engagement",
                      safe_get(digital, "engagement_assessment", default="—").replace("_", " ").title())
        platforms = digital.get("social_platforms", [])
        if platforms:
            add_heading_text(doc, "Social Platforms", level=3)
            for p in platforms:
                add_arrow_bullet(doc, p)

    # Recent Brand Activity
    activity = brand.get("recent_brand_activity", [])
    if activity:
        add_heading_text(doc, "Recent Brand Activity", level=2)
        for a in activity:
            add_text_with_citation(doc, f"{a.get('date', '—')}: {a.get('fact', '—')}")
            add_styled_paragraph(
                doc,
                f"Source: {a.get('source_name', '—')}",
                font_size=8, italic=True, color=DARK_CYAN
            )

    # Confidence
    conf = brand.get("research_confidence", "—")
    add_styled_paragraph(
        doc,
        f"Section Research Confidence: {conf.upper()}",
        font_size=9, italic=True, color=DARK_CYAN, space_before=8
    )

    doc.add_page_break()


def build_research_notes(doc, data):
    """Build the Research Notes & Citations section."""
    add_section_label(doc, "Research Notes & Citations")
    add_horizontal_rule(doc)

    # Cross-section insights
    insights = data.get("cross_section_insights", [])
    if insights:
        add_heading_text(doc, "Cross-Section Insights", level=2)
        for i, insight in enumerate(insights, 1):
            add_styled_paragraph(doc, f"{i}. {insight}", font_size=10,
                                 space_after=6)

    # Key gaps
    gaps = data.get("key_gaps", [])
    if gaps:
        add_heading_text(doc, "Key Research Gaps", level=2)
        for gap in gaps:
            add_arrow_bullet(doc, gap, color=GREY_MID)

    # Bibliography
    citations = data.get("citations", [])
    if citations:
        add_heading_text(doc, "Bibliography", level=2)
        add_styled_paragraph(
            doc,
            f"{len(citations)} sources cited in this research profile.",
            font_size=9, italic=True, color=DARK_CYAN, space_after=8
        )

        for cite in citations:
            cid = cite.get("id", "?")
            name = cite.get("source_name", "—")
            url = cite.get("url", "—")
            ctype = cite.get("type", "—").replace("_", " ").title()
            accessed = cite.get("accessed_date", "—")
            desc = cite.get("description", "")

            p = doc.add_paragraph()

            # Citation number
            num_run = p.add_run(f"[{cid}]  ")
            num_run.font.name = FONT_NAME
            num_run.font.size = Pt(9)
            num_run.font.bold = True
            num_run.font.color.rgb = LIME_GREEN

            # Source name
            name_run = p.add_run(f"{name}")
            name_run.font.name = FONT_NAME
            name_run.font.size = Pt(9)
            name_run.font.bold = True
            name_run.font.color.rgb = EERIE_BLACK

            # Type and date
            meta_run = p.add_run(f"  ({ctype} | Accessed: {accessed})")
            meta_run.font.name = FONT_NAME
            meta_run.font.size = Pt(8)
            meta_run.font.color.rgb = DARK_CYAN

            p.paragraph_format.space_after = Pt(1)

            # URL
            url_p = doc.add_paragraph()
            url_run = url_p.add_run(f"    {url}")
            url_run.font.name = FONT_NAME
            url_run.font.size = Pt(8)
            url_run.font.color.rgb = DARK_CYAN
            url_run.font.italic = True
            url_p.paragraph_format.space_after = Pt(1)

            # Description
            if desc:
                desc_p = doc.add_paragraph()
                desc_run = desc_p.add_run(f"    {desc}")
                desc_run.font.name = FONT_NAME
                desc_run.font.size = Pt(8)
                desc_run.font.color.rgb = GREY_MID
                desc_p.paragraph_format.space_after = Pt(6)

    # Data sources fallback (if no citations array but data_sources exist)
    elif data.get("data_sources"):
        add_heading_text(doc, "Data Sources", level=2)
        for src in data["data_sources"]:
            if isinstance(src, dict):
                add_arrow_bullet(doc, f"{src.get('name', '—')} — {src.get('url', '—')}")
            else:
                add_arrow_bullet(doc, str(src))


# ===================================================================
# DOCUMENT ASSEMBLY
# ===================================================================

def add_header_footer(doc, data):
    """Add document header and footer with branding."""
    section = doc.sections[0]

    # Header
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_para.add_run(
        f"{data.get('company_name', 'Brand Research')} — Brand Deep Research Profile"
    )
    run.font.name = FONT_NAME
    run.font.size = Pt(8)
    run.font.color.rgb = DARK_CYAN
    run.font.bold = True

    # Footer with page numbers
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add "Prepared by Upscale" text
    footer_run = footer_para.add_run("Prepared by Upscale  |  ")
    footer_run.font.name = FONT_NAME
    footer_run.font.size = Pt(7)
    footer_run.font.color.rgb = DARK_CYAN

    # Add page number field
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    footer_para._p.append(fldChar1)
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    footer_para._p.append(instrText)
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    footer_para._p.append(fldChar2)


def try_add_logo(doc):
    """Attempt to add the Upscale logo to the document."""
    # Try to find logo in common locations
    logo_paths = [
        "/tmp/upscale-brand-skill/assets/logos/full/Full_Colour.png",
        os.path.join(os.path.dirname(__file__), "..", "assets", "logos",
                     "Full_Colour.png"),
    ]

    # Try SVG → PNG conversion
    svg_paths = [
        "/tmp/upscale-brand-skill/assets/logos/full/Full_Colour.svg",
    ]

    # Check for existing PNG
    for path in logo_paths:
        if os.path.exists(path):
            try:
                doc.add_picture(path, width=Inches(2.0))
                return True
            except Exception:
                continue

    # Try SVG conversion
    for svg_path in svg_paths:
        if os.path.exists(svg_path):
            try:
                import cairosvg
                png_path = "/tmp/upscale_logo_temp.png"
                cairosvg.svg2png(url=svg_path, write_to=png_path,
                                output_width=800)
                doc.add_picture(png_path, width=Inches(2.0))
                return True
            except Exception:
                continue

    return False


def create_document(data, output_path):
    """Create the branded Word document from research data."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)   # A4
    section.page_height = Inches(11.69)  # A4
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(10)
    font.color.rgb = EERIE_BLACK

    # Add header/footer
    add_header_footer(doc, data)

    # Try to add logo
    logo_added = try_add_logo(doc)
    if logo_added:
        add_horizontal_rule(doc)

    # Build all sections
    build_cover_page(doc, data)
    build_company_overview(doc, data)
    build_products_services(doc, data)
    build_performance(doc, data)
    build_outlook_strategy(doc, data)
    build_brand_profile(doc, data)
    build_research_notes(doc, data)

    # Save
    doc.save(output_path)
    print(f"Word document created: {output_path}")


# ===================================================================
# CLI ENTRY POINT
# ===================================================================

def main():
    """Main entry point."""
    if len(sys.argv) < 2:
        print("Usage: python build_brand_research_docx.py <json_path> [output_path]")
        print("\nExample:")
        print("  python build_brand_research_docx.py brand_research.json output.docx")
        sys.exit(1)

    json_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else "brand_deep_research.docx"

    try:
        with open(json_path, "r") as f:
            data = json.load(f)
        create_document(data, output_path)
        print(f"Success! Brand research document saved to: {output_path}")
    except FileNotFoundError:
        print(f"Error: JSON file not found: {json_path}")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"Error: Invalid JSON in {json_path}: {e}")
        sys.exit(1)
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


if __name__ == "__main__":
    main()
